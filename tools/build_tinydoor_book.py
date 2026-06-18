#!/usr/bin/env python3
"""Build the TinyDoor book DOCX from canonical Markdown options.

This script intentionally treats the Markdown files in Options/ as the source of
truth. It strips license/nav boilerplate, preserves option ordering from
Options/INDEX.md, and creates a reproducible Word document with page breaks
between options.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable, Iterator, NamedTuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


class OptionEntry(NamedTuple):
    number: int
    label: str
    path: Path


NAV_RE = re.compile(
    r"<!--\s*TINYDOOR_NAV_START\s*-->.*?<!--\s*TINYDOOR_NAV_END\s*-->",
    re.DOTALL,
)
HTML_COMMENT_RE = re.compile(r"\A\s*<!--.*?-->\s*", re.DOTALL)
INDEX_LINK_RE = re.compile(r"^\s*(\d{3})\.\s+\[(.*?)\]\(([^)]+\.md)\)", re.MULTILINE)
INLINE_TOKEN_RE = re.compile(r"(\*\*_[^*]+?_\*\*|\*\*[^*]+?\*\*|_[^_]+?_|`[^`]+?`)")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build TinyDoor DOCX from Options Markdown")
    parser.add_argument("--repo-root", type=Path, default=Path.cwd())
    parser.add_argument("--index", type=Path, default=Path("Options/INDEX.md"))
    parser.add_argument("--through", type=int, default=40, help="Highest option number to include, e.g. 40")
    parser.add_argument("--output", type=Path, default=Path("options/book/tiny_door_and_pillow_forts.docx"))
    parser.add_argument("--title", default="Tiny Door and Pillow Forts")
    parser.add_argument("--author", default='Kevin "Andie" Williams')
    return parser.parse_args()


def read_index(repo_root: Path, index_path: Path, through: int) -> list[OptionEntry]:
    index_text = (repo_root / index_path).read_text(encoding="utf-8")
    entries: list[OptionEntry] = []
    for match in INDEX_LINK_RE.finditer(index_text):
        option_number = int(match.group(1))
        if option_number > through:
            continue
        entries.append(
            OptionEntry(
                number=option_number,
                label=match.group(2).strip(),
                path=(repo_root / "Options" / match.group(3)).resolve(),
            )
        )
    if not entries:
        raise ValueError(f"No options found in {index_path} through {through}")
    missing = [entry.path for entry in entries if not entry.path.exists()]
    if missing:
        missing_list = "\n".join(str(path) for path in missing)
        raise FileNotFoundError(f"Missing option Markdown files:\n{missing_list}")
    return entries


def clean_markdown(text: str) -> list[str]:
    text = text.replace("\ufeff", "")
    text = HTML_COMMENT_RE.sub("", text)
    text = NAV_RE.sub("", text)
    lines: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if line.strip() == "---":
            continue
        if not line.strip() and (not lines or not lines[-1].strip()):
            continue
        lines.append(line)
    while lines and not lines[-1].strip():
        lines.pop()
    return lines


def ensure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name in ["Heading 1", "Heading 2"]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.color.rgb = RGBColor(31, 78, 121)

    h1 = styles["Heading 1"]
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(12)

    h2 = styles["Heading 2"]
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    if "Option Label" not in styles:
        option_label = styles.add_style("Option Label", WD_STYLE_TYPE.PARAGRAPH)
    else:
        option_label = styles["Option Label"]
    option_label.font.name = "Aptos"
    option_label.font.size = Pt(9)
    option_label.font.bold = True
    option_label.font.color.rgb = RGBColor(89, 89, 89)
    option_label.paragraph_format.space_after = Pt(2)

    if "Tiny Quote" not in styles:
        quote = styles.add_style("Tiny Quote", WD_STYLE_TYPE.PARAGRAPH)
    else:
        quote = styles["Tiny Quote"]
    quote.font.name = "Aptos"
    quote.font.size = Pt(11)
    quote.font.italic = True
    quote.paragraph_format.left_indent = Inches(0.35)
    quote.paragraph_format.space_before = Pt(4)
    quote.paragraph_format.space_after = Pt(8)

    if "Tiny Code" not in styles:
        code = styles.add_style("Tiny Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Tiny Code"]
    code.font.name = "Courier New"
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.space_after = Pt(2)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_document(doc: Document, title: str, author: str, through: int) -> None:
    core = doc.core_properties
    core.title = title
    core.author = author
    core.subject = f"Options 001-{through:03d}"

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer.add_run("Tiny Door / You Have Another Option | Page ")
    add_page_number(footer)


def add_markdown_runs(paragraph, text: str) -> None:
    pos = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        run_text = token
        bold = False
        italic = False
        code = False
        if token.startswith("**_") and token.endswith("_**"):
            run_text = token[3:-3]
            bold = True
            italic = True
        elif token.startswith("**") and token.endswith("**"):
            run_text = token[2:-2]
            bold = True
        elif token.startswith("_") and token.endswith("_"):
            run_text = token[1:-1]
            italic = True
        elif token.startswith("`") and token.endswith("`"):
            run_text = token[1:-1]
            code = True
        run = paragraph.add_run(run_text)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_markdown_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    add_markdown_runs(paragraph, text)


def add_option(doc: Document, lines: list[str], *, is_first: bool) -> None:
    if not is_first:
        doc.add_page_break()

    in_code_block = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
        if not stripped:
            continue
        if in_code_block:
            doc.add_paragraph(line, style="Tiny Code")
            continue
        if stripped.startswith("# "):
            heading_text = stripped[2:].strip()
            if heading_text.lower().startswith("option "):
                doc.add_paragraph(heading_text, style="Option Label")
            else:
                doc.add_heading(heading_text, level=1)
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            continue
        if stripped.startswith(">"):
            add_markdown_paragraph(doc, stripped.lstrip(">").strip(), style="Tiny Quote")
            continue
        add_markdown_paragraph(doc, stripped)


def build_docx(repo_root: Path, entries: list[OptionEntry], output_path: Path, title: str, author: str, through: int) -> None:
    doc = Document()
    ensure_styles(doc)
    configure_document(doc, title, author, through)

    for idx, entry in enumerate(entries):
        lines = clean_markdown(entry.path.read_text(encoding="utf-8"))
        add_option(doc, lines, is_first=idx == 0)

    output_abs = (repo_root / output_path).resolve()
    output_abs.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_abs)


def main() -> None:
    args = parse_args()
    repo_root = args.repo_root.resolve()
    entries = read_index(repo_root, args.index, args.through)
    build_docx(repo_root, entries, args.output, args.title, args.author, args.through)
    print(f"Built {args.output} from {len(entries)} options through {args.through:03d}")


if __name__ == "__main__":
    main()
